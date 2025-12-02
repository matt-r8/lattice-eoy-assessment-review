#!/usr/bin/env python3
"""
lattice_reviews.py – fetch and summarise Lattice review-cycle feedback
"""

import os
import sys
import re
import html
import io
import contextlib
from pathlib import Path
from collections import defaultdict
import argparse
import logging
from functools import lru_cache
from io import StringIO
from docx import Document
from pathlib import Path
import io, logging, argparse
import requests
from dotenv import load_dotenv

# ─────────────────────────────  Logging  ──────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

# ─────────────────────────────  Env & session  ────────────────────────
load_dotenv()
BASE_URL = os.getenv("LATTICE_API_URL")
TOKEN    = os.getenv("LATTICE_API_TOKEN")

if not TOKEN or not BASE_URL:
    logger.error("Missing LATTICE_API_URL or LATTICE_API_TOKEN in environment.")
    sys.exit(1)

SESSION = requests.Session()
SESSION.headers.update({"Authorization": f"Bearer {TOKEN}"})

# ─────────────────────────────  Constants  ────────────────────────────
REVIEW_LIMIT   = 100        # per-reviewee page size
REVIEWEE_LIMIT = 100        # reviewees page size
CYCLE_LIMIT    = 1          # pull last cycle by default
TIMEOUT        = (30, 120)  # (connect, read)

# ─────────────────────────────  Helpers – API  ────────────────────────
def _get(url: str):
    resp = SESSION.get(url, timeout=TIMEOUT)
    resp.raise_for_status()
    return resp.json()

def get_me():
    return _get(f"{BASE_URL}/me")

def get_direct_reports(uid):
    return _get(f"{BASE_URL}/user/{uid}/directReports")["data"]

def get_all_review_cycles():
    url   = f"{BASE_URL}/reviewCycles?limit={CYCLE_LIMIT}"
    cycles = []
    while url:
        out   = _get(url)
        cycles.extend(out["data"])
        url = (
            f"{BASE_URL}/reviewCycles?limit={CYCLE_LIMIT}&startingAfter={out['endingCursor']}"
            if out.get("hasMore") else None
        )
    return cycles

@lru_cache(maxsize=1024)
def get_user_by_id(uid):
    return _get(f"{BASE_URL}/user/{uid}")

@lru_cache(maxsize=4096)
def get_question_by_id(qid):
    return _get(f"{BASE_URL}/question/{qid}")

@lru_cache(maxsize=4096)
def get_reviewee_by_id(rid):
    return _get(f"{BASE_URL}/reviewee/{rid}")

def get_reviewees_for_cycle(cid):
    url  = f"{BASE_URL}/reviewCycle/{cid}/reviewees?limit={REVIEWEE_LIMIT}"
    rvrs = []
    while url:
        out  = _get(url)
        rvrs.extend(out["data"])
        url = (
            f"{BASE_URL}/reviewCycle/{cid}/reviewees?limit={REVIEWEE_LIMIT}&startingAfter={out['endingCursor']}"
            if out.get("hasMore") else None
        )
    return rvrs

def get_reviews_for_reviewee(rid):
    url   = f"{BASE_URL}/reviewee/{rid}/reviews?limit={REVIEW_LIMIT}"
    reviews = []
    while url:
        out   = _get(url)
        reviews.extend(out["data"])
        url = (
            f"{BASE_URL}/reviewee/{rid}/reviews?limit={REVIEW_LIMIT}&startingAfter={out['endingCursor']}"
            if out.get("hasMore") else None
        )
    return reviews

# ─────────────────────────────  Fuzzy-key logic  ──────────────────────
def get_fuzzy_key(text: str) -> str:
    t = text.lower()
    if "growth mindset"        in t: return "growth"
    if "feedback"              in t: return "feedback"
    if "communicat"            in t: return "communication"
    if "embody"     in t or "values" in t: return "values"
    if "proactive"   in t or "anticipate risks" in t: return "proactive"
    if "hypothesis"           in t: return "hypothesis"
    if "collaborate" in t or "milestones" in t: return "collaboration"
    if "ownership"   in t or "responsibilities" in t: return "ownership"
    if "adapt"       in t or "adaptability"   in t: return "adaptability"
    if "rehire"                   in t: return "rehire"
    if "start, stop" in t or "continue doing" in t or "personal wins" in t:
        return "continue"
    return "other"

# ───────────────────────── summary printer ────────────────────────────
def print_summary(
    reviewee_id: str,
    cycle_id:    str,
    *,
    question_filter: str | None = None,
    group_by:        str        = "reviewer",
) -> None:
    """Print a review summary for one reviewee in one cycle."""

    # ── basic user info ────────────────────────────────────────────────
    user_id   = get_reviewee_by_id(reviewee_id)["user"]["id"]
    user_name = get_user_by_id(user_id)["name"]

    # pull **all** reviews for stats + later detail
    reviews = get_reviews_for_reviewee(reviewee_id)

    # ── overall averages (skip sentinel 6 = “no answer”) ───────────────
    self_scores:  list[float] = []
    peer_scores:  list[float] = []

    for r in reviews:
        if not r or not r.get("response"):
            continue
        val = r.get("response", {}).get("ratingString")
        try:
            score = float(val)
        except (TypeError, ValueError):
            continue
        if score == 6:                       # “insufficient experience” → ignore
            continue
        if r.get("reviewType") == "Self":
            self_scores.append(score)
        else:
            peer_scores.append(score)

    def _avg(lst: list[float]) -> float | None:
        return round(sum(lst) / len(lst), 2) if lst else None

    self_avg  = _avg(self_scores)
    peer_avg  = _avg(peer_scores)
    delta     = round(self_avg - peer_avg, 2) if self_avg and peer_avg else None

    logger.info(f"\n📋 Review Summary for: {user_name}")
    logger.info("Overall Scores")
    if peer_avg is not None:
        logger.info(f"- Peers: {peer_avg}")
    if self_avg is not None:
        logger.info(f"- Self:  {self_avg}")
    if delta is not None:
        logger.info(f"- Delta: {delta}")

    # ── build per-question / per-reviewer buckets ──────────────────────
    reviews_by_reviewer: defaultdict = defaultdict(list)
    question_cache: dict[str, dict]  = {}

    def _q_cached(qid: str) -> dict:
        if qid not in question_cache:
            question_cache[qid] = get_question_by_id(qid)
        return question_cache[qid]

    for review in reviews:
        if not review or not review.get("question"):
            continue
        reviewer = get_user_by_id(review["reviewer"]["id"])
        q        = _q_cached(review["question"]["id"])
        q_text   = html.unescape(q["body"])

        if question_filter:
            if re.escape(question_filter.lower()) not in q_text.lower():
                continue

        scale = {str(v["value"]): v["descriptor"] for v in q.get("ratingScale") or []}
        resp  = review.get("response") or {}

        rating_str  = resp.get("ratingString", "N/A")
        rating_desc = scale.get(rating_str, "Unknown")

        comment = resp.get("comment") or ""
        comment = html.unescape(comment)
        comment = re.sub(r"<br\s*/?>", "\n", comment).strip() or None

        r_type = review.get("reviewType", "Peer")
        reviews_by_reviewer[(reviewer["id"], reviewer["name"], r_type)].append(
            (q_text, rating_str, rating_desc, comment)
        )

    if not reviews_by_reviewer:
        logger.warning("No matching questions found.")
        return

    # ── output by reviewer ────────────────────────────────────────────
    if group_by == "reviewer":
        for (rid, name, r_type), resps in reviews_by_reviewer.items():
            logger.info(f"\n👤 Reviewer: {name} ({r_type})")
            for i, (q, val, desc, com) in enumerate(resps, 1):
                comment_str = f" 💬 {com}" if com else ""
                logger.info(f"Q{i}: {q}\nA: {val}, {desc} - {name} ({r_type}){comment_str}")

    # ── output by question ────────────────────────────────────────────
    else:  # group_by == "question"
        qmap: defaultdict = defaultdict(list)
        for (rid, name, r_type), resps in reviews_by_reviewer.items():
            for q, val, desc, com in resps:
                fkey = get_fuzzy_key(q)
                if fkey == "other":
                    continue
                qmap[fkey].append((q, name, rid, val, desc, com, r_type))

        for i, (fkey, resps) in enumerate(qmap.items(), 1):
            logger.info(f"\n🔸 {fkey.upper()} Question ({i})")
            logger.info(resps[0][0])  # question body once

            scores, self_score, self_comment = [], None, None
            peer_resps: list = []

            for _, name, _, val, _, com, r_type in resps:
                try:
                    score = float(val)
                except (TypeError, ValueError):
                    score = None

                if r_type == "Self":
                    if score != 6:
                        self_score = score
                    self_comment = com
                else:
                    if score not in (None, 6):
                        scores.append(score)
                    peer_resps.append((name, val, com))

            # self line
            if self_score is not None or self_comment:
                ss = f"{self_score}" if self_score is not None else "None"
                logger.info(f"- Self: {ss}{' 💬 '+self_comment if self_comment else ''}")

            # peer averages & delta
            if scores:
                p_avg = round(sum(scores) / len(scores), 2)
                logger.info(f"- Peers: {p_avg}")
                if self_score is not None:
                    logger.info(f"- 🔍 Score Delta: {round(self_score - p_avg, 2)}")

            # individual peer lines
            if peer_resps:
                logger.info("- Peer Responses:")
                for name, val, com in peer_resps:
                    logger.info(f"  - {name}: {val}{' 💬 '+com if com else ''}")

# ─────────────────────────────  CLI & main  ───────────────────────────
def print_question_help() -> None:
    logger.info(
        """
Available fuzzy question filters
--------------------------------
growth         • Growth mindset
feedback       • Feedback habits
communication  • Communication
values         • Company values
proactive      • Proactive behavior
hypothesis     • Hypothesis-led strategy
collaboration  • Goal collaboration
ownership      • Accountability
adaptability   • Adaptability
rehire         • Would re-hire?
continue       • Start / Stop / Continue
"""
    )

def find_user_by_name(name: str, users: list[dict]) -> dict | None:
    n = name.strip().lower()
    return next((u for u in users if u["name"].lower() == n), None)

def save_summary_to_docx(text: str, path: Path) -> None:
    """Write 〈text〉 (already newline-separated) to a Word .docx file."""
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    path = path.with_suffix(".docx")
    doc.save(path)

def extract_overall(text: str) -> tuple[float | None, float | None, float | None]:
    """
    Returns (peer_avg, self_avg, delta) as floats or None.

    Only the first Peers/Self/Delta lines that appear right after the
    'Overall Scores' header are considered, so per-question lines do not
    pollute the result.
    """
    # isolate the Overall-Scores stanza (three lines max)
    m_hdr = re.search(r"Overall Scores([\s\S]*?)^(?:🔸|👤|$)", text, re.M)
    stanza = m_hdr.group(1) if m_hdr else ""

    def _grab(label: str) -> float | None:
        m = re.search(rf"^- +{label}:\s+([-+]?\d+\.?\d*)", stanza, re.M)
        return float(m.group(1)) if m else None

    return _grab("Peers"), _grab("Self"), _grab("Delta")

def write_stack_rank(stack, path: Path) -> None:
    """stack -> list[(name, peer_avg, self_avg, delta)]  →  .docx table."""
    doc = Document()
    doc.add_heading("Direct-Report Stack Rank", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "Name"
    tbl.rows[0].cells[1].text = "Peers"
    tbl.rows[0].cells[2].text = "Self"
    tbl.rows[0].cells[3].text = "Δ"
    for name, peer, self_, delta in stack:
        row = tbl.add_row().cells
        row[0].text = name
        row[1].text = f"{peer:.2f}"  if peer  is not None else "—"
        row[2].text = f"{self_:.2f}" if self_ is not None else "—"
        row[3].text = f"{delta:+.2f}" if delta is not None else "—"
    doc.save(path.with_suffix(".docx"))

def save_summary_to_docx(text: str, path: Path) -> None:
    """
    Convert a newline-separated string to a Word document.
    """
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    doc.save(path)

# ------------------------------------------------------------------------
def main() -> None:
    ap = argparse.ArgumentParser(description="Fetch Lattice review summaries.")
    ap.add_argument("--cycle", default="2025 Formal Mid-Year Feedback",
                    help="Review-cycle name (exact match)")
    ap.add_argument("--filter", help="Fuzzy keyword to filter a single question")
    ap.add_argument("--group-by", choices=["reviewer", "question"],
                    default="reviewer", help="Group output (default: reviewer)")
    ap.add_argument("--help-questions", action="store_true",
                    help="Show available fuzzy keywords and exit")
    args = ap.parse_args()

    if args.help_questions:
        print_question_help(); return

    me       = get_me()
    reports  = get_direct_reports(me["id"])
    cycle    = next((c for c in get_all_review_cycles()
                     if c["name"] == args.cycle), None)
    if not cycle:
        logger.error("Review cycle not found."); return

    logger.info(f"Review Cycle : {cycle['name']}")
    logger.info(f"Manager      : {me['name']} ({me['id']})")
    logger.info(f"Direct Reports: {len(reports)} found\n" + "-" * 60)

    reviewees = get_reviewees_for_cycle(cycle["id"])       # one fetch

    out_dir = Path("summaries"); out_dir.mkdir(exist_ok=True)
    stack: list[tuple[str, float | None, float | None, float | None]] = []

    for rpt in reports:
        rvw = next((r for r in reviewees
                    if get_reviewee_by_id(r["id"])["user"]["id"] == rpt["id"]), None)
        if not rvw:
            logger.warning(f"No review data for {rpt['name']}"); continue

        # capture a single summary
        buf, handler = io.StringIO(), logging.StreamHandler(io.StringIO())
        handler.setFormatter(logging.Formatter("%(message)s"))
        logger.addHandler(handler); handler.stream = buf

        try:
            print_summary(rvw["id"], cycle["id"],
                          question_filter=args.filter, group_by=args.group_by)
        finally:
            logger.removeHandler(handler); handler.close()

        # write per-report docx
        save_summary_to_docx(buf.getvalue(), out_dir / rpt["name"].replace(" ", "_"))
        logger.info(f"↳ saved summary for {rpt['name']}")

        # collect for stack rank
        stack.append((rpt["name"], *extract_overall(buf.getvalue())))

    # build stack-rank file
    stack_sorted = sorted(stack, key=lambda t: (t[1] is None, -(t[1] or 0)))
    write_stack_rank(stack_sorted, out_dir / "_stack_rank.docx")
    logger.info("\n📄 _stack_rank.docx written to summaries/")

if __name__ == "__main__":
    main()