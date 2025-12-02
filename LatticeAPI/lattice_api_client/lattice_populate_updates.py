# Create a Drive

#!/usr/bin/env python3
"""
lattice_populate_updates.py
───────────────────────────
• Populate / refresh a local SQLite DB with Lattice “Updates”.
• Optional: restrict ingest to your own direct-reports.
• Optional: print updates for one author already stored in the DB.

DB file : lattice_updates.db
Table   : updates
"""

from __future__ import annotations
import os, sys, sqlite3, logging, re, argparse, shutil, textwrap
from typing import Any, Dict, Iterator
from functools import lru_cache
import html
import time, random, requests
from collections import defaultdict, Counter
from datetime import datetime, date, timedelta, timezone
from dotenv import load_dotenv
from docx import Document
import json

# ── logging ────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger(__name__)

# ── environment / session ──────────────────────────────────────────────
load_dotenv()
BASE = os.getenv("LATTICE_API_URL")
TOK  = os.getenv("LATTICE_API_TOKEN")
if not BASE or not TOK:
    log.error("Missing LATTICE_API_URL or LATTICE_API_TOKEN."); sys.exit(1)

SESSION = requests.Session()
SESSION.headers.update({"Authorization": f"Bearer {TOK}"})
TIMEOUT = 15

team_key = os.getenv("TEAM", "default").lower()
team_map_file = f"team_map_{team_key}.json" if team_key != "default" else "team_map.json"

with open(team_map_file) as f:
    TEAM_MAP = json.load(f)

# ── tiny HTTP helpers ──────────────────────────────────────────────────
RETRYABLE = {429, 502, 503, 504}

def _get(url: str, *, max_retry: int = 5) -> Dict[str, Any]:
    """GET with exponential back-off on 429 / 5xx."""
    for attempt in range(max_retry):
        resp = SESSION.get(url, timeout=TIMEOUT)
        if resp.status_code not in RETRYABLE:
            resp.raise_for_status()
            return resp.json()

        # hit rate-limit/server-busy
        wait = (resp.headers.get("Retry-After")
                or 1.5 * 2**attempt + random.uniform(0, 0.5))
        wait = float(wait)
        log.warning(f"{resp.status_code} — retrying in {wait:.1f}s …")
        time.sleep(wait)

    resp.raise_for_status()

@lru_cache(maxsize=1024)
def user_name(uid: str) -> str:
    return _get(f"{BASE}/user/{uid.strip()}")["name"]

def format_last_first(full_name: str) -> str:
    parts = full_name.strip().split(" ", 1)
    return f"{parts[1]}, {parts[0]}" if len(parts) == 2 else full_name

def get_me() -> Dict[str, Any]:
    return _get(f"{BASE}/me")

def get_direct_reports(uid: str) -> list[Dict[str, Any]]:
    return _get(f"{BASE}/user/{uid}/directReports")["data"]

# ── iterate over updates (all or page-size) ────────────────────────────
def all_updates(page_size: int = 100) -> Iterator[Dict[str, Any]]:
    url: str | None = f"{BASE}/updates?limit={page_size}"
    while url:
        blk = _get(url)
        yield from blk["data"]
        url = (f"{BASE}/updates?limit={page_size}&startingAfter={blk['endingCursor']}"
               if blk.get("hasMore") else None)

def get_update_by_id(upd_id: str) -> Dict[str, Any]:
    """Fetch a single update - used when the list endpoint had no text."""
    return _get(f"{BASE}/update/{upd_id}")

# ── SQLite boiler-plate ────────────────────────────────────────────────
DB_PATH = "lattice_updates.db"
SCHEMA  = """
CREATE TABLE IF NOT EXISTS updates (
  id               TEXT,
  author_id        TEXT,
  author_name      TEXT,
  created_at       TEXT,
  question         TEXT,
  answer           TEXT,
  sentiment_object TEXT,
  sentiment_rating INTEGER,
  PRIMARY KEY (id, question)
);
"""

def open_db() -> sqlite3.Connection:
    con = sqlite3.connect(DB_PATH)
    con.execute("PRAGMA journal_mode=WAL;")
    con.execute(SCHEMA)
    return con

def to_iso(ts: int) -> str:
    return datetime.fromtimestamp(ts, timezone.utc).date().isoformat()

def show_updates_for_name(full_name: str) -> None:
    con = open_db()
    con.row_factory = sqlite3.Row
    rows = con.execute(
        "SELECT id, created_at, question, answer, sentiment_object, sentiment_rating FROM updates "
        "WHERE author_name = ? ORDER BY created_at DESC, id, question",
        (full_name,)
    ).fetchall()

    grouped = {}
    for row in rows:
        key = (row["id"], row["created_at"], row["sentiment_object"], row["sentiment_rating"])
        grouped.setdefault(key, []).append((row["question"], row["answer"]))

    width = shutil.get_terminal_size(fallback=(120, 20)).columns
    log.info(f"\n📝 Updates for {full_name} ({len(grouped)})\n" + "-"*width)
    for (u_id, date, sent_obj, sent_rating), qas in grouped.items():
        log.info(f"{date} — Sentiment: {sent_rating if sent_rating is not None else 'N/A'}")
        for q, a in qas:
            log.info(textwrap.fill(q.strip(), width=width))
            log.info(textwrap.fill(a.strip(), width=width) + "\n")

def clean_text(raw: str) -> str:
    text = html.unescape(raw)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    return text.strip()

# ── export updates to DOCX ─────────────────────────────────────────────
def export_updates_to_docx(name: str, updates: list[tuple], team: str = "Unknown") -> None:
    doc = Document()
    doc.add_heading(f"Updates for {name}", level=1)

    # Group by month → then by date
    updates_by_month = defaultdict(lambda: defaultdict(list))
    ratings_by_date = {}

    for date, q, a, rating in updates:
        month = datetime.strptime(date, "%Y-%m-%d").strftime("%Y-%m")
        updates_by_month[month][date].append((q, a))
        if rating is not None:
            ratings_by_date[date] = rating

    months = sorted(updates_by_month.keys(), reverse=True)
    if len(months) >= 2:
        cur_month, prev_month = months[0], months[1]
        cur_scores = [ratings_by_date[d] for d in updates_by_month[cur_month] if d in ratings_by_date]
        prev_scores = [ratings_by_date[d] for d in updates_by_month[prev_month] if d in ratings_by_date]
        cur_avg = round(sum(cur_scores) / len(cur_scores), 2) if cur_scores else 0
        prev_avg = round(sum(prev_scores) / len(prev_scores), 2) if prev_scores else 0
        delta = cur_avg - prev_avg
        arrow = "⬆️" if delta > 0 else "⬇️" if delta < 0 else "➡️"
        doc.add_paragraph(
            f"📈 Trend: {cur_avg} {arrow} from last month ({prev_avg})",
            style="Intense Quote"
        )

    for month in sorted(updates_by_month.keys(), reverse=True):
        month_scores = [ratings_by_date[d] for d in updates_by_month[month] if d in ratings_by_date]
        if month_scores:
            avg = round(sum(month_scores) / len(month_scores), 2)
            doc.add_paragraph(
                f"📊 {month} Avg Sentiment: {avg} (from {len(month_scores)} entries)",
                style="Intense Quote"
            )

        for date in sorted(updates_by_month[month], reverse=True):
            doc.add_heading(str(date), level=2)
            rating = ratings_by_date.get(date)
            doc.add_paragraph(f"Sentiment: {rating if rating is not None else 'N/A'}", style="Intense Quote")

            for i, (q, a) in enumerate(updates_by_month[month][date], 1):
                doc.add_paragraph(f"Q{i}: {q}")
                para = doc.add_paragraph()
                para.add_run(f"A{i}: {a.strip().replace(chr(10), ' ')}")

    os.makedirs("updates", exist_ok=True)
    clean_name = name.rstrip(",")
    prefix = f"{team}_" if team != "Unknown" else ""
    fname = os.path.join("updates", f"{prefix}{clean_name}.docx")
    doc.save(fname)
    log.debug(f"Saved: {fname}")


def export_all_direct_reports() -> None:
    me = get_me()
    direct_reports = get_direct_reports(me["id"])
    con = open_db()
    team_sentiments = defaultdict(list)

    for dr in direct_reports:
        name = format_last_first(dr["name"])
        team = TEAM_MAP.get(name, "Unknown")
        if team == "Unknown":
            log.warning(f"Missing team mapping for: {name}")

        log.debug(f"Checking: {name}")
        first_last = " ".join(reversed(name.split(", ")))
        rows = con.execute(
            "SELECT created_at, question, answer, sentiment_rating FROM updates "
            "WHERE author_name = ? ORDER BY created_at DESC, question",
            (first_last,)
        ).fetchall()

        if not rows:
            log.warning(f"No updates found in DB for: {name}")
        else:
            export_updates_to_docx(name, rows, team=team)

        # team summary data
        rows = con.execute(
            "SELECT created_at, sentiment_rating FROM updates WHERE author_name = ?",
            (first_last,)
        ).fetchall()
        seen = set()
        for created_date, rating in rows:
            if rating is not None:
                month = datetime.strptime(created_date, "%Y-%m-%d").strftime("%Y-%m")
                key = (first_last, created_date)
                if key not in seen:
                    seen.add(key)
                    log.debug(f"[{team}] {first_last} → {created_date} = {rating}")
                    team_sentiments[(team, month)].append(rating)

    # regroup: month → team → ratings
    month_to_team_ratings = defaultdict(lambda: defaultdict(list))
    for (team, month), ratings in team_sentiments.items():
        month_to_team_ratings[month][team] = ratings
    log.debug("Month → Team → Entry Count")
    for month, teams in month_to_team_ratings.items():
        for team, ratings in teams.items():
            log.debug(f"{month} → {team}: {len(ratings)} entries")

    doc = Document()
    doc.add_heading("Team Sentiment Summary", level=1)

    # Compute org-wide monthly averages and percent delta
    all_months = sorted(month_to_team_ratings.keys(), reverse=True)
    if len(all_months) >= 2:
        cur_month, prev_month = all_months[0], all_months[1]
        cur_scores = [score for ratings in month_to_team_ratings[cur_month].values() for score in ratings]
        prev_scores = [score for ratings in month_to_team_ratings[prev_month].values() for score in ratings]
        cur_avg = round(sum(cur_scores) / len(cur_scores), 2) if cur_scores else 0
        prev_avg = round(sum(prev_scores) / len(prev_scores), 2) if prev_scores else 0
        delta = cur_avg - prev_avg
        trend = "⬆️" if delta > 0 else "⬇️" if delta < 0 else "➡️"
        doc.add_paragraph(
            f"📈 Practice Average: {cur_avg} ({trend} from last month {prev_avg})",
            style="Intense Quote"
        )

    for month in sorted(month_to_team_ratings.keys(), reverse=True):
        # 📈 Participation calc
        active_authors = con.execute("""
            SELECT COUNT(DISTINCT author_name) FROM updates
            WHERE strftime('%Y-%m', created_at) = ?
        """, (month,)).fetchone()[0]
        percent = round((active_authors / len(direct_reports)) * 100, 1)
        doc.add_paragraph(f"👍 {month} — Participation: {percent}% with 1+ updates", style="Intense Quote")

        sorted_teams = sorted(
            month_to_team_ratings[month].items(),
            key=lambda item: sum(item[1]) / len(item[1]) if item[1] else 0,
            reverse=True
        )
        for team, ratings in sorted_teams:
            avg = round(sum(ratings) / len(ratings), 2)
            doc.add_paragraph(
                f"{team}: {avg} (from {len(ratings)} {'Risers' if len(ratings) > 1 else 'Riser'})",
            )
        doc.add_paragraph()

    # Find latest week (last Sunday to Saturday)
    today = date.today()
    last_sunday = today - timedelta(days=today.weekday() + 1)
    this_week = last_sunday.strftime("%Y-%m-%d")

    # Get distinct authors from the last 7 days
    con = open_db()
    active_names = con.execute("""
        SELECT DISTINCT author_name FROM updates
        WHERE created_at >= date('now', '-7 days')
    """).fetchall()
    active_count = len(active_names)
    total_reports = len(get_direct_reports(get_me()["id"]))
    percent = round((active_count / total_reports) * 100, 1)
    
    os.makedirs("updates", exist_ok=True)
    doc.save("updates/_team_sentiment_summary.docx")
    log.debug("Saved: updates/_team_sentiment_summary.docx")

def export_one_report(name: str) -> None:
    con = open_db()
    rows = con.execute(
        "SELECT created_at, question, answer, sentiment_rating FROM updates "
        "WHERE author_name = ? ORDER BY created_at DESC, question",
        (name,)
    ).fetchall()
    if rows:
        export_updates_to_docx(name, rows)
    else:
        log.warning(f"No updates found for '{name}'.")

# ── main populate routine ──────────────────────────────────────────────
def populate_updates(page_size: int, restrict_to_directs: bool) -> None:
    me = get_me()
    direct_ids = {u["id"] for u in get_direct_reports(me["id"])} if restrict_to_directs else None
    if restrict_to_directs:
        log.info(f"Ingest restricted to {len(direct_ids)} direct reports.")

    con = open_db(); cur = con.cursor()
    ins, skip = 0, 0
    log.info("  Fetching updates …")

    for upd in all_updates(page_size):

        # 1️⃣  extract the author ID every time
        uid = ((upd.get("user")    or {}).get("id") or
            (upd.get("author")  or {}).get("id") or
            (upd.get("creator") or {}).get("id") or
            "unknown")
        print(f"Processing update {upd['id']} from user {uid}, Name: {user_name(uid)}")
        # 2️⃣  optional filtering to directs
        if restrict_to_directs and uid not in direct_ids:
            continue

        # 3️⃣  skip duplicates already in DB
        u_id = upd["id"]
        if cur.execute("SELECT 1 FROM updates WHERE id=?", (u_id,)).fetchone():
            skip += 1
            continue

        # 4️⃣  normal insert path
        try:
            full = get_update_by_id(u_id)
            raw_text = full.get("text")
            text = clean_text(raw_text) if isinstance(raw_text, str) else ""
            responses = []

            if not text and isinstance(full.get("responses"), dict):
                responses = full["responses"].get("data", [])
                parts = []
                for r in responses:
                    q = r.get("question")
                    a = r.get("answer")
                    q = q.strip() if isinstance(q, str) else ""
                    a = a.strip() if isinstance(a, str) else ""
                    if q or a:
                        parts.append(f"{q}\n{a}".strip())
                text = clean_text("\n\n".join(parts))

            sent = full.get("sentiment") or {}
        except Exception as e:
            log.warning(f"  Could not fetch full text for {u_id[:8]}: {e}")
            responses = []
            sent = {}

        sent_obj = sent.get("object")
        sent_rating = sent.get("rating")

        for r in responses:
            q = (r.get("question") or "").strip()
            a = (r.get("answer") or "").strip()
            if q or a:
                cur.execute(
                    "INSERT INTO updates (id, author_id, author_name, created_at, question, answer, sentiment_object, sentiment_rating) "
                    "VALUES (?,?,?,?,?,?,?,?)",
                    (u_id, uid, user_name(uid), to_iso(upd["createdAt"]), clean_text(q), clean_text(a), sent_obj, sent_rating)
                )
        ins += 1

    con.commit(); con.close()
    log.info(f"  Inserted {ins} new rows, skipped {skip} existing.")
    log.info(f"DB path: {os.path.abspath(DB_PATH)}")

# ── CLI glue ───────────────────────────────────────────────────────────
def main() -> None:
    ap = argparse.ArgumentParser(description="Sync Lattice Updates into local SQLite DB")
    ap.add_argument("--page", type=int, default=100)
    ap.add_argument("--directs", action="store_true")
    ap.add_argument("--report-name")
    ap.add_argument("--export-docx", nargs="?", const="ALL")
    args = ap.parse_args()

    if not args.report_name and not args.export_docx:
        populate_updates(args.page, args.directs)
    elif args.report_name:
        show_updates_for_name(args.report_name)
    elif args.export_docx == "ALL":
        export_all_direct_reports()
    elif args.export_docx:
        export_one_report(args.export_docx)

if __name__ == "__main__":
    start = time.time()
    main()
    end = time.time()
    log.info(f"⏱️  Completed in {end - start:.2f} seconds.")